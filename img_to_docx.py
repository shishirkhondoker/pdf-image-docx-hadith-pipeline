from docx import Document
from docx.shared import Inches
import os

# Path to the images folder
images_folder = 'images/'

# Path for the output Word document
output_docx = 'output_images.docx'

# Create a new Word document
doc = Document()

# Add a title or heading
doc.add_heading('Converted PDF Images', 0)

# Loop through all the images in the images folder
for filename in os.listdir(images_folder):
    if filename.endswith(".png"): 
        image_path = os.path.join(images_folder, filename)
        
        # Add the image to the Word document
        doc.add_paragraph(f"Image: {filename}")  
        doc.add_picture(image_path, width=Inches(6))  
        doc.add_paragraph()  

# Save the Word document
doc.save(output_docx)

print(f"Images have been successfully added to {output_docx}")